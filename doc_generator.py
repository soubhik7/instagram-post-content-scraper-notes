import os
import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ---------------------------------------------------------------------------
# Helper utilities
# ---------------------------------------------------------------------------

def _set_cell_bg(cell, hex_str: str):
    """Set the background shading of a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_str)
    tcPr.append(shd)


def _remove_table_borders(table):
    """Remove all visible borders from a table."""
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "none")
        el.set(qn("w:sz"), "0")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "auto")
        tblBorders.append(el)
    tblPr.append(tblBorders)


def _add_divider(doc):
    """Add a thin indigo horizontal rule paragraph."""
    para = doc.add_paragraph()
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "6366F1")
    pBdr.append(bottom)
    pPr.append(pBdr)
    para.paragraph_format.space_after = Pt(4)
    para.paragraph_format.space_before = Pt(0)
    return para


def _add_page_number(para):
    """Append a PAGE field code to an existing paragraph."""
    run = para.add_run()
    fldChar_begin = OxmlElement("w:fldChar")
    fldChar_begin.set(qn("w:fldCharType"), "begin")
    run._r.append(fldChar_begin)

    run2 = para.add_run()
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = " PAGE "
    run2._r.append(instrText)

    run3 = para.add_run()
    fldChar_end = OxmlElement("w:fldChar")
    fldChar_end.set(qn("w:fldCharType"), "end")
    run3._r.append(fldChar_end)


def _styled_run(para, text: str, size: int, bold: bool = False,
                italic: bool = False, color: str = None, font: str = "Calibri"):
    """Add a formatted run to a paragraph."""
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = font
    run.font.size = Pt(size)
    if color:
        r, g, b = int(color[0:2], 16), int(color[2:4], 16), int(color[4:6], 16)
        run.font.color.rgb = RGBColor(r, g, b)
    return run


def _set_cell_margins(cell, top=0, start=0, bottom=0, end=0):
    """Set inner cell margins in twips (1pt = 20 twips)."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for side, val in (("top", top), ("left", start), ("bottom", bottom), ("right", end)):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:w"), str(val))
        el.set(qn("w:type"), "dxa")
        tcMar.append(el)
    tcPr.append(tcMar)


# ---------------------------------------------------------------------------
# Main document generator
# ---------------------------------------------------------------------------

def generate_document(raw_dir: str) -> str:
    """
    Reads the downloaded .txt and image files from raw_dir, compiles them into
    a polished Word document, and saves it as <shortcode>_combined.docx in the
    same directory.  Returns the absolute path to the created docx.
    """
    # --- Discover source files --------------------------------------------------
    txt_file = None
    image_files = []

    for filename in os.listdir(raw_dir):
        if filename.endswith(".txt"):
            txt_file = os.path.join(raw_dir, filename)
        elif filename.lower().endswith((".jpg", ".jpeg", ".png")):
            image_files.append(os.path.join(raw_dir, filename))

    image_files.sort()

    # --- Read caption -----------------------------------------------------------
    if txt_file and os.path.exists(txt_file):
        with open(txt_file, "r", encoding="utf-8") as f:
            caption = f.read()
    else:
        caption = "(No text content found in this post)"

    # --- Create document --------------------------------------------------------
    doc = Document()

    # Page setup: A4, 1.8 cm margins all sides
    section = doc.sections[0]
    section.page_width  = Cm(21)
    section.page_height = Cm(29.7)
    for attr in ("left_margin", "right_margin", "top_margin", "bottom_margin"):
        setattr(section, attr, Cm(1.8))

    # ---- 1. Header banner ------------------------------------------------------
    banner_table = doc.add_table(rows=1, cols=1)
    _remove_table_borders(banner_table)
    banner_cell = banner_table.cell(0, 0)
    _set_cell_bg(banner_cell, "312E81")
    _set_cell_margins(banner_cell, top=240, start=300, bottom=200, end=300)

    # Clear default empty paragraph inside the cell
    for p in banner_cell.paragraphs:
        p._element.getparent().remove(p._element)

    # Title + date on one line
    title_para = banner_cell.add_paragraph()
    title_para.paragraph_format.space_before = Pt(0)
    title_para.paragraph_format.space_after  = Pt(0)
    _styled_run(title_para, "TechNotes", size=24, bold=True,
                color="FFFFFF", font="Calibri Light")
    today = datetime.datetime.now().strftime("%B %d, %Y")
    _styled_run(title_para, f"   |   {today}", size=10, italic=True,
                color="C7D2FE", font="Calibri Light")

    # ---- 3. Caption section ----------------------------------------------------
    def add_section_heading(doc, text: str):
        para = doc.add_paragraph()
        para.paragraph_format.space_before = Pt(6)
        para.paragraph_format.space_after  = Pt(1)
        _styled_run(para, text, size=12, bold=True,
                    color="6366F1", font="Calibri Light")
        _add_divider(doc)

    add_section_heading(doc, "Post Caption")

    caption_lines = caption.split("\n")
    for line in caption_lines:
        if line.strip():
            para = doc.add_paragraph()
            para.paragraph_format.space_after  = Pt(3)
            para.paragraph_format.line_spacing = Pt(14)
            _styled_run(para, line, size=10, color="1E1B4B", font="Calibri")
        else:
            blank = doc.add_paragraph()
            blank.paragraph_format.space_after = Pt(0)

    # ---- 4. Images section -----------------------------------------------------
    if image_files:
        add_section_heading(doc, "Images")

        # Lay images out in a 2-column table so multiple fit per page
        img_table = doc.add_table(rows=0, cols=2)
        _remove_table_borders(img_table)

        for i in range(0, len(image_files), 2):
            row = img_table.add_row()
            for col, img_path in enumerate(image_files[i:i+2]):
                cell = row.cells[col]
                _set_cell_margins(cell, top=80, start=80, bottom=80, end=80)
                para = cell.paragraphs[0]
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                para.paragraph_format.space_before = Pt(0)
                para.paragraph_format.space_after  = Pt(0)
                try:
                    para.add_run().add_picture(img_path, width=Inches(3.0))
                except Exception as exc:
                    _styled_run(para,
                                f"[Could not insert: {os.path.basename(img_path)} — {exc}]",
                                size=8, italic=True, color="DC2626", font="Calibri")

    # ---- 5. Footer -------------------------------------------------------------
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_para.paragraph_format.space_before = Pt(0)
    footer_para.paragraph_format.space_after  = Pt(0)

    _styled_run(footer_para, "TechNotes  •  Page ",
                size=9, color="94A3B8", font="Calibri")
    _add_page_number(footer_para)
    _styled_run(footer_para, f"  •  {datetime.date.today().year}",
                size=9, color="94A3B8", font="Calibri")

    # ---- Save ------------------------------------------------------------------
    timestamp = datetime.datetime.now().strftime("%d_%B_%Y_%I%M%p")
    doc_path = os.path.join(raw_dir, f"Notes_{timestamp}.docx")
    doc.save(doc_path)
    for f in os.listdir(raw_dir):
        fp = os.path.join(raw_dir, f)
        if os.path.isfile(fp) and fp != doc_path:
            os.remove(fp)
    return os.path.abspath(doc_path)


if __name__ == "__main__":
    import sys
    if len(sys.argv) > 1:
        res = generate_document(sys.argv[1])
        print("Document generated at:", res)
    else:
        print("Usage: python doc_generator.py <path_to_raw_dir>")
